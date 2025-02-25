import os
import zipfile
import random
import re
import string
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime, timezone
import tkinter as tk
from tkinter import filedialog
from bs4 import BeautifulSoup
from docx.shared import Pt

#Generate the current time
current_time_iso = datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')

print("Select the file you want to convert")
root = tk.Tk()
input_path = filedialog.askopenfilename(title='Select .docx or .celtx file for converting', filetypes=[("Word Documents, Celtx files", "*.docx *.celtx"), ("Word Documents", "*.docx"), ("Celtx files", "*.celtx")])
root.withdraw()

screenplay_filename = os.path.basename(input_path)
screenplay_title = screenplay_filename.split('.')[0]
print(f"Screenplay Title: {screenplay_title}")
file_format = screenplay_filename.split('.')[1]
if file_format == 'docx':
  print("Converting .docx to .celtx")
  celtx_path = input_path.replace('.docx', '.celtx')

  # Generate ids
  def generate_id(length):
      chars = string.ascii_letters + string.digits
      return ''.join(random.choice(chars) for _ in range(length))

  project_id = generate_id(12)
  project_code = generate_id(3)

  #Define functions
  def generate_scene_id():
      # Generate 5 random chars (letters/numbers) + 1 number + 00
      chars = string.ascii_letters + string.digits
      random_part = ''.join(random.choice(chars) for _ in range(5))
      number_part = str(random.randint(0, 9))
      return f"{random_part}{number_part}00"

  screenplay_head_text = r"""<!DOCTYPE HTML PUBLIC "-//W3C//DTD HTML 4.01//EN" "http://www.w3.org/TR/html4/strict.dtd">
  <html>
  <head>
    <title>"""+ screenplay_title +""""</title>
    <link rel="stylesheet" type="text/css" href="chrome://celtx/content/editor.css">
    <meta http-equiv="Content-Type" content="text/html; charset=UTF-8">
    <style id="leftheading" type="text/css">p.sceneheading:before { display: none !important; } </style>
    <style id="rightheading" type="text/css">p.sceneheading:after { display: none !important; } </style>
    <meta content="none" name="CX.sceneNumbering">
    <meta content="false" name="CX.showPageNumbers">
    <meta content="false" name="CX.showFirstPageNumber">
    <meta content="false" name="CX.showCharNumbers">
    <meta content="false" name="CX.dialogNumbering">
    <style id="pagenumbers" type="text/css">.softbreak { display: none !important; } </style>
    <style id="charnumbers" type="text/css">
  .character:before, .sound:before, .music:before, .voice:before { display: none !important; }

    </style>
    <link href="chrome://celtx/content/style/film/USLetter/Normal.css" type="text/css" rel="stylesheet">
    <meta content="AUTHOR NAME HERE" name="Author">
    <meta content="" name="DC.source">
    <meta content="" name="DC.rights">
    <meta content="" name="CX.contact">
    <meta content="By" name="CX.byline">
  </head>
  <body>
  """

  doc = Document(input_path)
  html_body = ""
  scene_count = 1

  #Setup character dictoinary, scene heading list and current_scene_id for project.rdf
  characters = {}
  scene_headings = []
  current_scene_id = ''

  #Parse the .docx file

  for paragraph in doc.paragraphs:
      if not paragraph.text.strip():
          continue
          
      text = paragraph.text.strip()
      alignment = paragraph.alignment
      
      # Handle scene headings
      if '.' in text:
          before_dot = text.split('.')[0]
          if before_dot.isupper() and before_dot:
              scene_id = generate_scene_id()
              current_scene_id = generate_scene_id()
              html_body += f'<p scenestr="{scene_count}" scenenumber="{scene_count}" id="{scene_id}" class="sceneheading">{text}<br>\n</p>\n'
              scene_headings.append((scene_count, scene_id, text, current_scene_id))
              scene_count += 1
              continue
      
      # Handle other elements
      if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER and text.isupper():
          element_class = "character"
          if text in characters.keys():
              scene_list = characters.get(text)
              if current_scene_id not in scene_list:
                  scene_list.append(current_scene_id)
                  characters[text] = scene_list
          else:
              characters[text] = [current_scene_id]
      elif alignment == WD_PARAGRAPH_ALIGNMENT.CENTER and text.startswith('(') and text.endswith(')'):
          element_class = "parenthetical"
      elif alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
          element_class = "dialog"
      elif alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT and text.isupper():
          element_class = "transition"
      elif text.isupper():
          element_class = "shot"
      else:
          element_class = "action"
          
      html_body += f'<p class="{element_class}">{text}<br>\n</p>\n'

      html_end = "</body>\n</html>"
  complete_html = screenplay_head_text + html_body + html_end

  # Generate local.rdf and scratch-###.html

  local_rdf_text = f"""<?xml version="1.0"?>
  <RDF:RDF xmlns:dc="http://purl.org/dc/elements/1.1/"
          xmlns:cx="http://celtx.com/NS/v1/"
          xmlns:rdfs="http://www.w3.org/2000/01/rdf-schema#"
          xmlns:NC="http://home.netscape.com/NC-rdf#"
          xmlns:RDF="http://www.w3.org/1999/02/22-rdf-syntax-ns#">
    <RDF:Description RDF:about="http://celtx.com/project/{project_id}">
      <cx:opentabs RDF:resource="rdf:#$3DD1u"/>
    </RDF:Description>
    <RDF:Seq RDF:about="rdf:#$3DD1u">
      <RDF:li RDF:resource="http://celtx.com/res/1kEqvojZWTHA"/>
    </RDF:Seq>
  </RDF:RDF>
  """
  scratchpad_text = r"""<!DOCTYPE HTML PUBLIC "-//W3C//DTD HTML 4.01//EN" "http://www.w3.org/TR/html4/strict.dtd">
  <html>
  <head>
    <title></title>
    <style id="scratch" type="text/css">
  @media screen {
    body { background-color: #FFFFCC !important; }
  }
      </style>
    <link rel="stylesheet" type="text/css" href="chrome://celtx/content/editor.css">
    <meta http-equiv="Content-Type" content="text/html; charset=UTF-8">
    <style id="leftheading" type="text/css">p.sceneheading:before { display: none !important; } </style>
    <style id="rightheading" type="text/css">p.sceneheading:after { display: none !important; } </style>
    <meta content="none" name="CX.sceneNumbering">
    <meta content="false" name="CX.showPageNumbers">
    <meta content="false" name="CX.showFirstPageNumber">
    <meta content="false" name="CX.showCharNumbers">
    <meta content="false" name="CX.dialogNumbering">
    <style id="pagenumbers" type="text/css">.softbreak { display: none !important; } </style>
    <style id="charnumbers" type="text/css">
  .character:before, .sound:before, .music:before, .voice:before { display: none !important; }

    </style>
    <link href="chrome://celtx/content/style/film/USLetter/Normal.css" type="text/css" rel="stylesheet">
  </head>
  <body>
  <p class="action"><br>
  </p>
  </body>
  </html>"""

  #Generate project.rdf
  project_rdf_head = f"""<?xml version="1.0"?>
  <RDF:RDF xmlns:dc="http://purl.org/dc/elements/1.1/"
          xmlns:cx="http://celtx.com/NS/v1/"
          xmlns:rdfs="http://www.w3.org/2000/01/rdf-schema#"
          xmlns:NC="http://home.netscape.com/NC-rdf#"
          xmlns:RDF="http://www.w3.org/1999/02/22-rdf-syntax-ns#">
    <RDF:Seq RDF:about="rdf:#$06OTd3">
    </RDF:Seq>
    <RDF:Seq RDF:about="rdf:#$W5OTd3">
      <RDF:li RDF:resource="rdf:#$.5OTd3"/>
    </RDF:Seq>
    <RDF:Seq RDF:about="rdf:#$Y5OTd3">
      <RDF:li RDF:resource="rdf:#$Z5OTd3"/>
      <RDF:li RDF:resource="rdf:#$+5OTd3"/>
    </RDF:Seq>
    <RDF:Seq RDF:about="http://celtx.com/res/9uEellKm1xdh">
      <RDF:li RDF:resource="http://celtx.com/res/InouvsBJ47GI"/>
      <RDF:li RDF:resource="http://celtx.com/res/1kEqvojZWTHA"/>
      <RDF:li RDF:resource="http://celtx.com/res/defaultsketchdoc"/>
    </RDF:Seq>
    <RDF:Description RDF:about="http://celtx.com/res/9uEellKm1xdh"
                    cx:projectRoot="true"
                    dc:title="Script" />
    <cx:Document RDF:about="http://celtx.com/res/InouvsBJ47GI"
                    dc:title="Master Catalog"
                    cx:size="2">
      <cx:doctype RDF:resource="http://celtx.com/NS/v1/CatalogDocument"/>
      <cx:filter RDF:resource="celtx:filter:all"/>
      <cx:members RDF:resource="rdf:#$U5OTd3"/>
      <cx:departments RDF:resource="rdf:#$Y5OTd3"/>
    </cx:Document>
    <RDF:Seq RDF:about="http://celtx.com/res/U0xHoW9oCmfD">
      <RDF:li RDF:resource="http://celtx.com/res/9uEellKm1xdh"/>
    </RDF:Seq>
    <RDF:Seq RDF:about="rdf:#$16OTd3">
    </RDF:Seq>
    <RDF:Seq RDF:about="rdf:#$X5OTd3">
    </RDF:Seq>
    <RDF:Seq RDF:about="rdf:#$U5OTd3">
      <RDF:li RDF:resource="http://celtx.com/res/r5hrsWgDPrNn"/>
      <RDF:li RDF:resource="http://celtx.com/res/uCdfBtKoUfXx"/>
    </RDF:Seq>
    <cx:Document RDF:about="http://celtx.com/res/1kEqvojZWTHA"
                    dc:title="{screenplay_title}"
                    cx:sidebarvisible="sidebarnotes"
                    cx:localFile="script-{project_code}.html"
                    cx:auxFile="scratch-{project_code}.html">
      <cx:doctype RDF:resource="http://celtx.com/NS/v1/ScriptDocument"/>
      <cx:scenes RDF:resource="rdf:#$V5OTd3"/>
      <cx:tagnames RDF:resource="rdf:#$26OTd3"/>
    </cx:Document>
  """

  project_rdf_characters = ''
  for i in characters.keys():
      character_id = generate_id(12)
      project_rdf_characters += f"""<cx:Cast RDF:about="http://celtx.com/res/{character_id}"
                    dc:title="{i}" />
  """
      seq_id = generate_id(6)
      project_rdf_characters += f"""<RDF:Seq RDF:about="rdf:#${seq_id}">
    """
      for i in characters[i]:
        project_rdf_characters += f"""<RDF:li RDF:resource="http://celtx.com/res/{i}"/>
    """
      project_rdf_characters = project_rdf_characters[:-2]
      project_rdf_characters += """</RDF:Seq>
  """
      
  project_rdf_scenes = ''
  for i in scene_headings:
    padded_ordinal = str(i[0]).zfill(4)
    parsed1 = i[2].split('. ')
    try:
        intext = parsed1[0]
        setting = parsed1[1]
        setting_split  = setting.split(' - ')
        place = setting_split[0]
        time = setting_split[1]
        project_rdf_scenes += f""""  <RDF:Description RDF:about="http://celtx.com/res/{i[3]}"
                      cx:sceneid="{i[1]}"
                      cx:location=" "
                      cx:ordinal="{str(i[0])}"
                      cx:sortord="{padded_ordinal}"
                      cx:intext="{intext}"
                      cx:setting="{setting}"
                      cx:daynight="{time}"
                      dc:title="{i[2]}">
        <cx:members RDF:resource="rdf:#$.1CWh3"/>
        <cx:markup RDF:resource="rdf:#$+1CWh3"/>
      </RDF:Description>"""
    except IndexError:
      intext = parsed1[0]
      setting = parsed1[1]
      project_rdf_scenes += f"""<RDF:Description RDF:about="http://celtx.com/res/apkKa400"
                    cx:sceneid="{i[1]}"
                    cx:location=" "
                    cx:ordinal="{str(i[0])}"
                    cx:sortord="{padded_ordinal}"
                    cx:intext="{intext}"
                    cx:setting="{setting}"
                    dc:title="{i[2]}">
      <cx:members RDF:resource="rdf:#$06OTd3"/>
      <cx:markup RDF:resource="rdf:#$16OTd3"/>
      </RDF:Description>"""

  project_rdf_end = f"""  <RDF:Seq RDF:about="rdf:#$.5OTd3">
      <RDF:li RDF:resource="http://celtx.com/res/r5hrsWgDPrNn"/>
      <RDF:li RDF:resource="http://celtx.com/res/uCdfBtKoUfXx"/>
    </RDF:Seq>
    <cx:DepartmentList RDF:about="rdf:#$.5OTd3"
                    cx:size="2">
      <cx:department RDF:resource="http://celtx.com/NS/v1/Cast"/>
    </cx:DepartmentList>
    <RDF:Seq RDF:about="rdf:#$26OTd3">
      <RDF:li>Plot A</RDF:li>
      <RDF:li>Plot B</RDF:li>
      <RDF:li>Plot C</RDF:li>
      <RDF:li>Plot D</RDF:li>
      <RDF:li>Plot E</RDF:li>
      <RDF:li>Plot F</RDF:li>
      <RDF:li>Plot G</RDF:li>
    </RDF:Seq>
    <cx:Project RDF:about="http://celtx.com/project/{project_id}"
                    cx:fileVersion="1.4"
                    dc:title="{screenplay_title}"
                    dc:modified="{current_time_iso}">
      <cx:components RDF:resource="http://celtx.com/res/U0xHoW9oCmfD"/>
    </cx:Project>
    <RDF:Seq RDF:about="rdf:#$Z5OTd3">
      <RDF:li RDF:resource="http://celtx.com/res/r5hrsWgDPrNn"/>
    </RDF:Seq>
    <cx:DepartmentList RDF:about="rdf:#$Z5OTd3">
      <cx:department RDF:resource="http://celtx.com/NS/v1/Cast"/>
    </cx:DepartmentList>
  </RDF:RDF>"""

  project_rdf = project_rdf_head + project_rdf_characters + project_rdf_scenes + project_rdf_end

  with zipfile.ZipFile(celtx_path, 'w') as zipf:
      zipf.writestr('local.rdf', local_rdf_text)
      zipf.writestr(f'scratch-{project_code}.html', scratchpad_text)
      zipf.writestr(f'script-{project_code}.html', complete_html)
      zipf.writestr('project.rdf', project_rdf)
      zipf.close()

  print(f"Succesfully created {screenplay_title}.celtx")
elif file_format == 'celtx':
  print("Converting .celtx to .docx")
  docx_path = input_path.replace('.celtx', '.docx')
  
  def apply_formatting(paragraph, class_name):
    run = paragraph.runs[0]
    run.font.name = "Courier New"
    if class_name in ["sceneheading", "character", "transition", "shot"]:
        run.text = run.text.upper()
    if class_name == "character":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif class_name == "parenthetical":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif class_name == "dialog":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif class_name == "transition":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run.font.size = Pt(12)

  def extract_html_from_zip(zip_path):
      with zipfile.ZipFile(zip_path, 'r') as zip_file:
          for filename in zip_file.namelist():
              if re.match(r'script-\w+\.html', filename):
                  with zip_file.open(filename) as file:
                      return file.read().decode("utf-8")
      return None

  zip_path = input_path
  html_content = extract_html_from_zip(zip_path)
  soup = BeautifulSoup(html_content, "html.parser")
  doc = Document()

  for p in soup.find_all("p"):
      class_name = p.get("class", [None])[0]  # Get first class if multiple
      text = p.get_text(strip=True)
      if class_name and text:
          paragraph = doc.add_paragraph(text)
          apply_formatting(paragraph, class_name)

  doc.save(docx_path)
  print(f"Succesfully created {screenplay_title}.docx")